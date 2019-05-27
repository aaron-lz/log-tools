#coding:gbk
import docx,os,datetime
from docx import Document


now = datetime.datetime.now()
textFile = "{0}.txt".format(now.strftime("%Y-%m-%d"))
with open(textFile,'w') as fp:pass

def GetDocx(FileName,textFile):

	
	document = Document(FileName)
	tables = document.tables
	Name = FileName.split("_")[1]
	with open(textFile,'a') as fp:
		fp.write("# _ _ _ _ _ _ _ _ _ _   Start[{0}]  _ _ _ _ _ _ _ _ _ _\n".format(Name))
		for table in tables:
			for row in table.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						if paragraph.text == '':continue
						fp.write(paragraph.text + '\n')
		fp.write("# ____________________   Stop[{0}]   ____________________\n\n".format(Name))		

		
if __name__ == '__main__':
	for filename in os.listdir(os.getcwd()):
		if filename.endswith('.docx'):
			GetDocx(filename,textFile)